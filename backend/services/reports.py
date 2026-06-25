"""Sinh báo cáo Word (.docx) và Excel (.xlsx) từ kết quả đánh giá."""
from __future__ import annotations
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import Workbook


def build_summary_docx(
    package: dict[str, Any], vendor_names: dict[int, str],
    evals: dict[int, Any], ranking: list[dict[str, Any]], out_path: Path,
) -> Path:
    doc = Document()
    doc.add_heading("BÁO CÁO TỔNG HỢP ĐÁNH GIÁ HSDT", level=0)
    doc.add_paragraph(f"Mã gói thầu: {package.get('ma_so')}")
    doc.add_paragraph(f"Tên gói thầu: {package.get('ten')}")

    doc.add_heading("Bảng xếp hạng", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
        "Hạng", "Nhà thầu", "Giá đánh giá", "Điểm KT")
    for r in ranking:
        cells = table.add_row().cells
        cells[0].text = str(r.get("rank") or "Không hợp lệ")
        cells[1].text = vendor_names.get(r["vendor_id"], str(r["vendor_id"]))
        cells[2].text = f"{r['evaluated_price']:,}"
        cells[3].text = f"{r['technical_score']:.1f}"

    for vid, ev in evals.items():
        doc.add_heading(f"Nhà thầu: {vendor_names.get(vid, vid)}", level=1)
        for group, label in [("legality", "Hợp lệ"), ("capacity", "Năng lực"),
                             ("technical", "Kỹ thuật")]:
            results = ev.get(group, [])
            if not results:
                continue
            doc.add_heading(label, level=2)
            for item in results:
                doc.add_paragraph(
                    f"- {item['criteria_ten']}: {item['result']} "
                    f"(điểm {item['score']}) | Dẫn chứng: {item['evidence']} "
                    f"[trang {item['page_ref']}]"
                )
        fin = ev["financial"]
        doc.add_heading("Tài chính", level=2)
        doc.add_paragraph(f"Giá đánh giá: {fin['evaluated_price']:,}")
        if fin["errors"]:
            doc.add_paragraph(f"Số lỗi số học phát hiện: {len(fin['errors'])}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def build_summary_xlsx(
    vendor_names: dict[int, str], evals: dict[int, Any],
    ranking: list[dict[str, Any]], out_path: Path,
) -> Path:
    wb = Workbook()
    wb.remove(wb.active)

    for sheet_name, group in [("Hop le", "legality"), ("Nang luc", "capacity"),
                              ("Ky thuat", "technical")]:
        ws = wb.create_sheet(sheet_name)
        ws.append(["Nhà thầu", "Tiêu chí", "Kết quả", "Điểm", "Dẫn chứng", "Trang"])
        for vid, ev in evals.items():
            for item in ev.get(group, []):
                ws.append([
                    vendor_names.get(vid, vid), item["criteria_ten"], item["result"],
                    item["score"], item["evidence"], str(item["page_ref"]),
                ])

    ws_fin = wb.create_sheet("Tai chinh")
    ws_fin.append(["Nhà thầu", "Giá đánh giá", "Số lỗi số học"])
    for vid, ev in evals.items():
        fin = ev["financial"]
        ws_fin.append([vendor_names.get(vid, vid), float(fin["evaluated_price"]),
                       len(fin["errors"])])

    ws_rank = wb.create_sheet("Xep hang")
    ws_rank.append(["Hạng", "Nhà thầu", "Giá đánh giá", "Điểm KT", "Hợp lệ"])
    for r in ranking:
        ws_rank.append([
            r.get("rank") or "-", vendor_names.get(r["vendor_id"], r["vendor_id"]),
            float(r["evaluated_price"]), round(r["technical_score"], 1),
            "Có" if r["eligible"] else "Không",
        ])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))
    return out_path
